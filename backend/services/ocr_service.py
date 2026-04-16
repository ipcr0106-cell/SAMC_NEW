"""
SAMC 수입식품 검역 AI — OCR / 텍스트 추출 서비스

파일 유형별 텍스트 추출 전략:
  - PDF       → PyMuPDF(fitz) 텍스트 추출 + fallback: 이미지 페이지는 Vision OCR
  - 이미지     → base64 인코딩 후 Vision API로 텍스트 인식
  - HWP/HWPX → parser-service (kordoc, Node.js) HTTP 호출
  - Excel     → openpyxl로 셀 데이터 읽기

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
⚠️  임시 전환 안내 (TEMPORARY — 개발/테스트용):
    Vision OCR은 현재 OpenAI(gpt-4o)를 사용합니다.
    최종 통합 단계에서는 반드시 Claude Vision으로 롤백할 것.
    (검색 키워드: "# >>> OPENAI TEMP")
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
"""

from __future__ import annotations

import base64
import io
import logging
import os
from pathlib import Path
from typing import Optional

import httpx

logger = logging.getLogger(__name__)

PARSER_SERVICE_URL = os.getenv("F0_PARSER_SERVICE_URL", "http://localhost:3001")
PARSER_SERVICE_TOKEN = os.getenv("F0_PARSER_SERVICE_TOKEN", "")


async def extract_text_from_file(
    file_bytes: bytes,
    file_name: str,
    mime_type: str,
) -> str:
    """파일 바이트에서 Raw 텍스트를 추출하는 디스패처.

    Args:
        file_bytes: 파일 바이너리 데이터
        file_name: 원본 파일명 (확장자 판별용)
        mime_type: MIME 타입

    Returns:
        추출된 Raw 텍스트 문자열
    """
    ext = Path(file_name).suffix.lower()
    logger.info(f"텍스트 추출 시작: {file_name} (ext={ext}, mime={mime_type}, size={len(file_bytes)}bytes)")

    text = ""
    if ext == ".pdf":
        text = await _extract_from_pdf(file_bytes)
    elif ext in (".hwp", ".hwpx"):
        text = await _extract_from_hwp(file_bytes, file_name)
    elif ext in (".xlsx", ".xls"):
        text = _extract_from_excel(file_bytes)
    elif ext == ".docx":
        text = _extract_from_docx(file_bytes)
    elif ext in (".png", ".jpg", ".jpeg", ".webp"):
        text = await _extract_from_image(file_bytes, mime_type)
    else:
        logger.warning(f"지원하지 않는 파일 형식: {ext} ({file_name})")
        return ""

    if text:
        logger.info(f"텍스트 추출 성공: {file_name} → {len(text)}자")
    else:
        logger.warning(f"텍스트 추출 결과 없음: {file_name}")

    return text


# ─────────────────────────────────────────────
# PDF 추출 (PyMuPDF)
# ─────────────────────────────────────────────

async def _extract_from_pdf(file_bytes: bytes) -> str:
    """PyMuPDF로 텍스트 추출. 텍스트 레이어가 없는 스캔 PDF는 이미지 fallback."""
    try:
        import fitz  # PyMuPDF
    except ImportError:
        logger.error("PyMuPDF(fitz)가 설치되지 않았습니다. pip install PyMuPDF")
        return ""

    doc = fitz.open(stream=file_bytes, filetype="pdf")
    pages_text: list[str] = []
    image_pages: list[int] = []

    for page_num in range(len(doc)):
        page = doc[page_num]
        text = page.get_text("text").strip()
        if text:
            pages_text.append(f"--- 페이지 {page_num + 1} ---\n{text}")
        else:
            # 텍스트 레이어 없음 → 이미지로 변환하여 OCR 대기열에 추가
            image_pages.append(page_num)

    # 이미지 페이지가 있으면 Claude Vision으로 OCR
    for page_num in image_pages:
        page = doc[page_num]
        pix = page.get_pixmap(dpi=200)
        img_bytes = pix.tobytes("png")
        ocr_text = await _extract_from_image(img_bytes, "image/png")
        if ocr_text:
            pages_text.append(f"--- 페이지 {page_num + 1} (OCR) ---\n{ocr_text}")

    doc.close()
    return "\n\n".join(pages_text)


# ─────────────────────────────────────────────
# HWP/HWPX 추출 (parser-service 경유)
# ─────────────────────────────────────────────

async def _extract_from_hwp(file_bytes: bytes, file_name: str) -> str:
    """parser-service (Node.js + kordoc)에 HTTP 요청하여 HWP 텍스트 추출."""
    url = f"{PARSER_SERVICE_URL}/parse"
    headers = {}
    if PARSER_SERVICE_TOKEN:
        headers["Authorization"] = f"Bearer {PARSER_SERVICE_TOKEN}"

    try:
        async with httpx.AsyncClient(timeout=60.0) as client:
            files = {"file": (file_name, io.BytesIO(file_bytes), "application/octet-stream")}
            response = await client.post(url, files=files, headers=headers)
            response.raise_for_status()
            data = response.json()
            return data.get("text", "")
    except httpx.HTTPStatusError as e:
        logger.error(f"parser-service 응답 오류: {e.response.status_code}")
        return ""
    except httpx.ConnectError:
        logger.error(f"parser-service 연결 실패: {PARSER_SERVICE_URL}")
        return ""


# ─────────────────────────────────────────────
# Excel 추출 (openpyxl)
# ─────────────────────────────────────────────

def _extract_from_excel(file_bytes: bytes) -> str:
    """openpyxl로 엑셀 전체 시트를 텍스트로 변환."""
    try:
        from openpyxl import load_workbook
    except ImportError:
        logger.error("openpyxl이 설치되지 않았습니다. pip install openpyxl")
        return ""

    wb = load_workbook(filename=io.BytesIO(file_bytes), read_only=True, data_only=True)
    all_text: list[str] = []

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows: list[str] = [f"[시트: {sheet_name}]"]
        for row in ws.iter_rows(values_only=True):
            cells = [str(cell) if cell is not None else "" for cell in row]
            if any(cells):
                rows.append(" | ".join(cells))
        all_text.append("\n".join(rows))

    wb.close()
    return "\n\n".join(all_text)


# ─────────────────────────────────────────────
# DOCX 추출 (python-docx)
# ─────────────────────────────────────────────

def _extract_from_docx(file_bytes: bytes) -> str:
    """python-docx로 Word 문서의 본문 + 표 텍스트 추출."""
    try:
        from docx import Document
    except ImportError:
        logger.error("python-docx가 설치되지 않았습니다. pip install python-docx")
        return ""

    try:
        doc = Document(io.BytesIO(file_bytes))
    except Exception as e:
        logger.error(f"DOCX 파일 열기 실패: {e}")
        return ""

    parts: list[str] = []

    # 본문 단락
    for para in doc.paragraphs:
        txt = (para.text or "").strip()
        if txt:
            parts.append(txt)

    # 표
    for t_idx, table in enumerate(doc.tables, start=1):
        parts.append(f"[표 {t_idx}]")
        for row in table.rows:
            cells = [(cell.text or "").strip() for cell in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))

    return "\n".join(parts)


# ─────────────────────────────────────────────
# 이미지 OCR (Vision API)
# ─────────────────────────────────────────────

_OCR_PROMPT = (
    "이 이미지에 포함된 모든 텍스트를 정확하게 추출해주세요. "
    "표가 있다면 Markdown 표 형태로 변환해주세요. "
    "언어는 원문 그대로 유지하되, 읽기 어려운 부분은 [불명확]로 표시해주세요."
)


async def _extract_from_image(image_bytes: bytes, mime_type: str) -> str:
    """이미지 Vision OCR 디스패처.

    ⚠️ 현재(임시): OpenAI gpt-4o Vision 사용
    ⚠️ 최종(복원): Claude Vision (_extract_from_image_claude) 로 교체
    """
    # 지원 MIME 타입 보정
    media_type = mime_type
    if media_type not in ("image/png", "image/jpeg", "image/webp", "image/gif"):
        media_type = "image/png"

    # >>> OPENAI TEMP — 최종 통합 시 _extract_from_image_claude 로 교체
    return await _extract_from_image_openai(image_bytes, media_type)
    # return await _extract_from_image_claude(image_bytes, media_type)
    # <<< OPENAI TEMP


# ─────────────────────────────────────────────
# >>> OPENAI TEMP — 최종 통합 시 제거 가능
# ─────────────────────────────────────────────

async def _extract_from_image_openai(image_bytes: bytes, media_type: str) -> str:
    """OpenAI gpt-4o Vision으로 이미지 텍스트 추출 — 개발/테스트용 임시 구현."""
    openai_api_key = os.getenv("F0_OPENAI_API_KEY", "")
    if not openai_api_key:
        logger.error("F0_OPENAI_API_KEY가 설정되지 않았습니다.")
        return ""

    try:
        from openai import AsyncOpenAI
    except ImportError:
        logger.error("openai 패키지가 설치되지 않았습니다. pip install openai")
        return ""

    b64_data = base64.b64encode(image_bytes).decode("utf-8")
    data_url = f"data:{media_type};base64,{b64_data}"
    model = os.getenv("F0_OPENAI_MODEL", "gpt-4o")

    try:
        client = AsyncOpenAI(api_key=openai_api_key)
        completion = await client.chat.completions.create(
            model=model,
            max_tokens=4096,
            messages=[
                {
                    "role": "user",
                    "content": [
                        {"type": "text", "text": _OCR_PROMPT},
                        {"type": "image_url", "image_url": {"url": data_url}},
                    ],
                }
            ],
        )
        return completion.choices[0].message.content or ""
    except Exception as e:
        logger.error(f"OpenAI Vision OCR 실패: {e}")
        return ""

# <<< OPENAI TEMP END


# ─────────────────────────────────────────────
# --- CLAUDE ORIGINAL (최종 통합 시 사용) ---
# ─────────────────────────────────────────────

async def _extract_from_image_claude(image_bytes: bytes, media_type: str) -> str:
    """Claude Vision으로 이미지 텍스트 추출 — 최종 프로덕션용."""
    anthropic_api_key = os.getenv("F0_ANTHROPIC_API_KEY", "")
    if not anthropic_api_key:
        logger.error("F0_ANTHROPIC_API_KEY가 설정되지 않았습니다.")
        return ""

    b64_data = base64.b64encode(image_bytes).decode("utf-8")

    try:
        import anthropic

        client = anthropic.AsyncAnthropic(api_key=anthropic_api_key)
        message = await client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=4096,
            messages=[
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "image",
                            "source": {
                                "type": "base64",
                                "media_type": media_type,
                                "data": b64_data,
                            },
                        },
                        {"type": "text", "text": _OCR_PROMPT},
                    ],
                }
            ],
        )
        return message.content[0].text
    except Exception as e:
        logger.error(f"Claude Vision OCR 실패: {e}")
        return ""
