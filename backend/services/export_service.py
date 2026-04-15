"""
SAMC — OCR 분석 결과를 DOCX / PDF로 내보내는 서비스.

포함 섹션:
  1. 기본 정보
  2. 원재료 배합비율
  3. 제조공정 (공정 코드 + 선정 근거 + 원문)
  4. 수출국 라벨 분석 (디자인 설명, 라벨 문구, 경고)
  5. 라벨 제품 이미지 (Vision 추출 이미지 + 텍스트 필드)

- DOCX: python-docx (한국어 완벽 지원)
- PDF: reportlab + CID 기본 폰트(HYSMyeongJo-Medium) — 외부 폰트 파일 불필요
"""

from __future__ import annotations

import io
import logging
from datetime import datetime
from typing import Any

logger = logging.getLogger(__name__)

# 라벨 이미지 텍스트 필드 레이블
_IMG_FIELD_LABELS = [
    ("label_product_name",   "제품명"),
    ("label_ingredients",    "원재료"),
    ("label_content_volume", "내용량"),
    ("label_origin",         "원산지"),
    ("label_manufacturer",   "제조사"),
    ("label_case_number",    "케이스 넘버"),
]


# ─────────────────────────────────────────────
# DOCX 생성
# ─────────────────────────────────────────────

def build_docx(
    parsed: dict[str, Any],
    *,
    product_name: str = "",
    case_id: str = "",
    label_images: list[dict] | None = None,
) -> bytes:
    from docx import Document
    from docx.shared import Pt, Cm, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    label_images = label_images or []
    doc = Document()

    # 기본 폰트 (한국어 안전)
    style = doc.styles["Normal"]
    style.font.name = "맑은 고딕"
    style.font.size = Pt(10)

    basic = parsed.get("basic_info", {}) or {}
    pname = basic.get("product_name") or product_name or ""
    ings  = parsed.get("ingredients", []) or []
    proc  = parsed.get("process_info", {}) or {}
    label = parsed.get("label_info", {}) or {}

    # ── 표지 헤더 ──
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run("SAMC 수입식품 OCR 분석 결과")
    run.bold = True
    run.font.size = Pt(18)

    if pname:
        sub_p = doc.add_paragraph()
        sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_run = sub_p.add_run(pname)
        sub_run.font.size = Pt(13)
        sub_run.font.color.rgb = RGBColor(0x0f, 0x17, 0x2a)

    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_p.add_run(
        f"생성일: {datetime.now().strftime('%Y-%m-%d %H:%M')}"
        + (f"   |   Case ID: {case_id}" if case_id else "")
    ).font.size = Pt(9)

    doc.add_paragraph()

    # ── 1. 기본 정보 ──
    doc.add_heading("1. 기본 정보", level=1)
    t = doc.add_table(rows=0, cols=2)
    t.style = "Light Grid Accent 1"
    for k, v in [
        ("제품명",        basic.get("product_name") or product_name or "-"),
        ("수출국",        basic.get("export_country") or "-"),
        ("최초 수입 여부", "예" if basic.get("is_first_import") else "아니오"),
        ("유기인증",      "예" if basic.get("is_organic") else "아니오"),
        ("OEM",          "예" if basic.get("is_oem") else "아니오"),
    ]:
        row = t.add_row().cells
        row[0].text = k
        row[1].text = str(v)
    doc.add_paragraph()

    # ── 2. 원재료 배합비율 ──
    doc.add_heading("2. 원재료 배합비율", level=1)
    if ings:
        tbl = doc.add_table(rows=1, cols=5)
        tbl.style = "Light Grid Accent 1"
        hdr = tbl.rows[0].cells
        for i, h in enumerate(["성분명", "비율(%)", "원산지", "INS", "CAS"]):
            hdr[i].text = h
            for p in hdr[i].paragraphs:
                for r in p.runs:
                    r.bold = True
        for ing in ings:
            r = tbl.add_row().cells
            r[0].text = str(ing.get("name") or "")
            r[1].text = str(ing.get("ratio") or "")
            r[2].text = str(ing.get("origin") or "")
            r[3].text = str(ing.get("ins_number") or "")
            r[4].text = str(ing.get("cas_number") or "")
    else:
        doc.add_paragraph("(추출된 원재료 없음)")
    doc.add_paragraph()

    # ── 3. 제조공정 ──
    doc.add_heading("3. 제조공정", level=1)
    codes   = proc.get("process_codes") or []
    reasons = proc.get("process_code_reasons") or []
    raw     = proc.get("raw_process_text") or ""

    p = doc.add_paragraph()
    p.add_run("공정 코드: ").bold = True
    p.add_run(", ".join(codes) if codes else "(없음)")

    # 공정 코드별 선정 근거
    if reasons:
        doc.add_paragraph("코드별 선정 근거:").runs[0].bold = True
        for r in reasons:
            code   = r.get("code") or ""
            reason = r.get("reason") or ""
            if code or reason:
                bullet = doc.add_paragraph(style="List Bullet")
                bullet.add_run(f"{code}  ").bold = True
                bullet.add_run(reason)

    if raw:
        doc.add_paragraph("OCR 추출 원문:").runs[0].bold = True
        for line in raw.splitlines():
            doc.add_paragraph(line or " ")
    doc.add_paragraph()

    # ── 4. 수출국 라벨 분석 ──
    doc.add_heading("4. 수출국 라벨 분석", level=1)
    ltexts = label.get("label_texts") or []
    warns  = label.get("warnings") or []
    desc   = label.get("design_description") or ""

    if desc:
        doc.add_paragraph("디자인 설명:").runs[0].bold = True
        doc.add_paragraph(desc)

    if ltexts:
        doc.add_paragraph("라벨 문구:").runs[0].bold = True
        for lt in ltexts:
            doc.add_paragraph(str(lt), style="List Bullet")

    if warns:
        doc.add_paragraph("경고/주의사항:").runs[0].bold = True
        for w in warns:
            doc.add_paragraph(str(w), style="List Bullet")

    if not (ltexts or warns or desc):
        doc.add_paragraph("(라벨 정보 없음)")
    doc.add_paragraph()

    # ── 5. 라벨 제품 이미지 ──
    if label_images:
        doc.add_heading("5. 라벨 제품 이미지", level=1)
        doc.add_paragraph(
            f"Vision AI가 자동 추출한 제품 이미지 {len(label_images)}개"
        ).runs[0].font.size = Pt(9)

        for img_data in label_images:
            img_bytes = img_data.get("bytes")
            idx = img_data.get("image_index", 0)

            doc.add_heading(f"이미지 {idx + 1}", level=2)

            # 이미지 삽입
            if img_bytes:
                try:
                    doc.add_picture(io.BytesIO(img_bytes), width=Cm(8))
                except Exception as e:
                    logger.warning(f"DOCX 이미지 삽입 실패 (idx={idx}): {e}")
                    doc.add_paragraph("(이미지 삽입 실패)")

            # 추출 텍스트 필드 테이블
            fields = [(lbl, img_data.get(key)) for key, lbl in _IMG_FIELD_LABELS if img_data.get(key)]
            if fields:
                ft = doc.add_table(rows=0, cols=2)
                ft.style = "Light Grid Accent 1"
                for lbl, val in fields:
                    row = ft.add_row().cells
                    row[0].text = lbl
                    row[0].paragraphs[0].runs[0].bold = True
                    row[1].text = str(val)
            else:
                doc.add_paragraph("(추출된 텍스트 없음)").runs[0].font.size = Pt(9)

            doc.add_paragraph()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ─────────────────────────────────────────────
# PDF 생성 (reportlab + 한국어 CID 폰트)
# ─────────────────────────────────────────────

_KOREAN_FONT_REGISTERED = False


def _register_korean_font():
    global _KOREAN_FONT_REGISTERED
    if _KOREAN_FONT_REGISTERED:
        return
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont
    pdfmetrics.registerFont(UnicodeCIDFont("HYSMyeongJo-Medium"))
    pdfmetrics.registerFont(UnicodeCIDFont("HYGothic-Medium"))
    _KOREAN_FONT_REGISTERED = True


def build_pdf(
    parsed: dict[str, Any],
    *,
    product_name: str = "",
    case_id: str = "",
    label_images: list[dict] | None = None,
) -> bytes:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
        Image as RLImage, KeepTogether,
    )

    label_images = label_images or []
    _register_korean_font()

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        leftMargin=1.8 * cm, rightMargin=1.8 * cm,
        topMargin=1.8 * cm, bottomMargin=1.8 * cm,
        title="SAMC OCR 분석 결과",
    )

    styles = getSampleStyleSheet()
    base  = ParagraphStyle("KR",      parent=styles["Normal"], fontName="HYSMyeongJo-Medium", fontSize=10, leading=14)
    h1    = ParagraphStyle("KR-H1",   parent=base, fontName="HYGothic-Medium",  fontSize=14, leading=18, spaceBefore=12, spaceAfter=6,  textColor=colors.HexColor("#0f172a"))
    h2    = ParagraphStyle("KR-H2",   parent=base, fontName="HYGothic-Medium",  fontSize=11, leading=15, spaceBefore=8,  spaceAfter=4,  textColor=colors.HexColor("#334155"))
    ttl   = ParagraphStyle("KR-Title",parent=base, fontName="HYGothic-Medium",  fontSize=18, leading=22, alignment=1, spaceAfter=8)
    sub   = ParagraphStyle("KR-Sub",  parent=base, fontName="HYGothic-Medium",  fontSize=13, leading=17, alignment=1, spaceAfter=4)
    meta  = ParagraphStyle("KR-Meta", parent=base, fontSize=9,  leading=12, alignment=1, textColor=colors.grey, spaceAfter=18)
    small = ParagraphStyle("KR-Sm",   parent=base, fontSize=9,  leading=12, textColor=colors.HexColor("#64748b"))
    bul   = ParagraphStyle("KR-Bul",  parent=base, leftIndent=12, fontSize=10, leading=14)

    basic  = parsed.get("basic_info", {}) or {}
    pname  = basic.get("product_name") or product_name or ""
    ings   = parsed.get("ingredients", []) or []
    proc   = parsed.get("process_info", {}) or {}
    label  = parsed.get("label_info", {}) or {}

    elems = []

    # ── 표지 헤더 ──
    elems.append(Paragraph("SAMC 수입식품 OCR 분석 결과", ttl))
    if pname:
        elems.append(Paragraph(pname, sub))
    meta_txt = f"생성일: {datetime.now().strftime('%Y-%m-%d %H:%M')}"
    if case_id:
        meta_txt += f"   |   Case ID: {case_id}"
    elems.append(Paragraph(meta_txt, meta))

    # ── 1. 기본 정보 ──
    elems.append(Paragraph("1. 기본 정보", h1))
    kv = [
        ["제품명",        basic.get("product_name") or product_name or "-"],
        ["수출국",        basic.get("export_country") or "-"],
        ["최초 수입 여부", "예" if basic.get("is_first_import") else "아니오"],
        ["유기인증",      "예" if basic.get("is_organic") else "아니오"],
        ["OEM",          "예" if basic.get("is_oem") else "아니오"],
    ]
    t = Table(kv, colWidths=[4 * cm, 12 * cm])
    t.setStyle(TableStyle([
        ("FONT",          (0, 0), (-1, -1), "HYSMyeongJo-Medium", 10),
        ("BACKGROUND",    (0, 0), (0, -1),  colors.HexColor("#f1f5f9")),
        ("GRID",          (0, 0), (-1, -1), 0.25, colors.HexColor("#cbd5e1")),
        ("VALIGN",        (0, 0), (-1, -1), "MIDDLE"),
        ("LEFTPADDING",   (0, 0), (-1, -1), 8),
        ("RIGHTPADDING",  (0, 0), (-1, -1), 8),
        ("TOPPADDING",    (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
    ]))
    elems.append(t)
    elems.append(Spacer(1, 12))

    # ── 2. 원재료 배합비율 ──
    elems.append(Paragraph("2. 원재료 배합비율", h1))
    if ings:
        data = [["성분명", "비율(%)", "원산지", "INS", "CAS"]]
        for ing in ings:
            data.append([
                str(ing.get("name") or ""),
                str(ing.get("ratio") or ""),
                str(ing.get("origin") or ""),
                str(ing.get("ins_number") or ""),
                str(ing.get("cas_number") or ""),
            ])
        t = Table(data, colWidths=[5.5*cm, 2*cm, 3*cm, 2.5*cm, 3*cm], repeatRows=1)
        t.setStyle(TableStyle([
            ("FONT",          (0, 0), (-1, -1), "HYSMyeongJo-Medium", 9),
            ("FONT",          (0, 0), (-1, 0),  "HYGothic-Medium", 9),
            ("BACKGROUND",    (0, 0), (-1, 0),  colors.HexColor("#0f172a")),
            ("TEXTCOLOR",     (0, 0), (-1, 0),  colors.white),
            ("GRID",          (0, 0), (-1, -1), 0.25, colors.HexColor("#cbd5e1")),
            ("ROWBACKGROUNDS",(0, 1), (-1, -1), [colors.white, colors.HexColor("#f8fafc")]),
            ("VALIGN",        (0, 0), (-1, -1), "MIDDLE"),
            ("LEFTPADDING",   (0, 0), (-1, -1), 5),
            ("RIGHTPADDING",  (0, 0), (-1, -1), 5),
            ("TOPPADDING",    (0, 0), (-1, -1), 4),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ]))
        elems.append(t)
    else:
        elems.append(Paragraph("(추출된 원재료 없음)", base))
    elems.append(Spacer(1, 12))

    # ── 3. 제조공정 ──
    elems.append(Paragraph("3. 제조공정", h1))
    codes   = proc.get("process_codes") or []
    reasons = proc.get("process_code_reasons") or []
    raw     = proc.get("raw_process_text") or ""

    elems.append(Paragraph(f"<b>공정 코드:</b> {', '.join(codes) if codes else '(없음)'}", base))

    if reasons:
        elems.append(Spacer(1, 6))
        elems.append(Paragraph("<b>코드별 선정 근거:</b>", base))
        for r in reasons:
            code   = (r.get("code") or "").replace("&", "&amp;").replace("<", "&lt;")
            reason = (r.get("reason") or "").replace("&", "&amp;").replace("<", "&lt;")
            if code or reason:
                elems.append(Paragraph(f"• <b>{code}</b>  {reason}", bul))

    if raw:
        elems.append(Spacer(1, 6))
        elems.append(Paragraph("<b>OCR 추출 원문:</b>", base))
        for line in raw.splitlines():
            safe = line.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            elems.append(Paragraph(safe or "&nbsp;", base))
    elems.append(Spacer(1, 12))

    # ── 4. 수출국 라벨 분석 ──
    elems.append(Paragraph("4. 수출국 라벨 분석", h1))
    ltexts   = label.get("label_texts") or []
    warns    = label.get("warnings") or []
    desc     = label.get("design_description") or ""
    any_label = False

    if desc:
        any_label = True
        elems.append(Paragraph("<b>디자인 설명:</b>", base))
        elems.append(Paragraph(desc.replace("<", "&lt;"), base))
        elems.append(Spacer(1, 6))
    if ltexts:
        any_label = True
        elems.append(Paragraph("<b>라벨 문구:</b>", base))
        for lt in ltexts:
            elems.append(Paragraph(f"• {str(lt).replace('<', '&lt;')}", bul))
        elems.append(Spacer(1, 6))
    if warns:
        any_label = True
        elems.append(Paragraph("<b>경고/주의사항:</b>", base))
        for w in warns:
            elems.append(Paragraph(f"• {str(w).replace('<', '&lt;')}", bul))
    if not any_label:
        elems.append(Paragraph("(라벨 정보 없음)", base))
    elems.append(Spacer(1, 12))

    # ── 5. 라벨 제품 이미지 ──
    if label_images:
        elems.append(Paragraph("5. 라벨 제품 이미지", h1))
        elems.append(Paragraph(
            f"Vision AI가 자동 추출한 제품 이미지 {len(label_images)}개",
            small,
        ))
        elems.append(Spacer(1, 8))

        for img_data in label_images:
            img_bytes = img_data.get("bytes")
            idx = img_data.get("image_index", 0)

            block = []
            block.append(Paragraph(f"이미지 {idx + 1}", h2))

            # 이미지
            if img_bytes:
                try:
                    rl_img = RLImage(io.BytesIO(img_bytes), width=7*cm, height=7*cm)
                    rl_img.hAlign = "LEFT"
                    block.append(rl_img)
                except Exception as e:
                    logger.warning(f"PDF 이미지 삽입 실패 (idx={idx}): {e}")
                    block.append(Paragraph("(이미지 삽입 실패)", small))

            # 추출 텍스트 필드
            fields = [(lbl, img_data.get(key)) for key, lbl in _IMG_FIELD_LABELS if img_data.get(key)]
            if fields:
                fd = [[lbl, str(val)] for lbl, val in fields]
                ft = Table(fd, colWidths=[3*cm, 13*cm])
                ft.setStyle(TableStyle([
                    ("FONT",         (0, 0), (-1, -1), "HYSMyeongJo-Medium", 9),
                    ("FONT",         (0, 0), (0, -1),  "HYGothic-Medium", 9),
                    ("BACKGROUND",   (0, 0), (0, -1),  colors.HexColor("#f1f5f9")),
                    ("GRID",         (0, 0), (-1, -1), 0.25, colors.HexColor("#cbd5e1")),
                    ("VALIGN",       (0, 0), (-1, -1), "TOP"),
                    ("LEFTPADDING",  (0, 0), (-1, -1), 6),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                    ("TOPPADDING",   (0, 0), (-1, -1), 4),
                    ("BOTTOMPADDING",(0, 0), (-1, -1), 4),
                ]))
                block.append(Spacer(1, 4))
                block.append(ft)
            else:
                block.append(Paragraph("(추출된 텍스트 없음)", small))

            elems.append(KeepTogether(block))
            elems.append(Spacer(1, 12))

    doc.build(elems)
    return buf.getvalue()
